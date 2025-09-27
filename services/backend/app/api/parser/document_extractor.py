import os
import uuid
from datetime import datetime
from typing import Optional
from pathlib import Path
from .models import LegalDocument, DocumentType


class DocumentExtractor:
    """Handles text extraction from various document formats."""
    
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
    
    def extract_text(self, file_path: str, document_type: DocumentType) -> str:
        """Extract text content from a document file."""
        try:
            if document_type == DocumentType.TXT:
                return self._extract_from_txt(file_path)
            elif document_type == DocumentType.PDF:
                return self._extract_from_pdf(file_path)
            elif document_type == DocumentType.DOCX:
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported document type: {document_type}")
        except Exception as e:
            raise Exception(f"Error extracting text from {file_path}: {str(e)}")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        try:
            import PyPDF2
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except ImportError:
            # Fallback to basic text extraction if PyPDF2 not available
            return self._fallback_pdf_extraction(file_path)
        except Exception as e:
            raise Exception(f"PDF extraction failed: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except ImportError:
            # Fallback if python-docx not available
            return self._fallback_docx_extraction(file_path)
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")
    
    def _fallback_pdf_extraction(self, file_path: str) -> str:
        """Fallback PDF extraction method."""
        # This is a placeholder - in a real implementation, you might use
        # other libraries like pdfplumber or pymupdf
        return f"[PDF content extraction not available for {file_path}]"
    
    def _fallback_docx_extraction(self, file_path: str) -> str:
        """Fallback DOCX extraction method."""
        # This is a placeholder - in a real implementation, you might use
        # other libraries or convert to text first
        return f"[DOCX content extraction not available for {file_path}]"
    
    def save_uploaded_file(self, file_content: bytes, filename: str) -> tuple[str, int]:
        """Save uploaded file and return file path and size."""
        # Generate unique filename to avoid conflicts
        file_extension = Path(filename).suffix
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        file_path = self.upload_dir / unique_filename
        
        # Save file
        with open(file_path, 'wb') as f:
            f.write(file_content)
        
        file_size = len(file_content)
        return str(file_path), file_size
    
    def get_document_type(self, filename: str) -> DocumentType:
        """Determine document type from filename."""
        extension = Path(filename).suffix.lower()
        
        if extension == '.txt':
            return DocumentType.TXT
        elif extension == '.pdf':
            return DocumentType.PDF
        elif extension in ['.docx', '.doc']:
            return DocumentType.DOCX
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    
    def create_document_record(self, filename: str, content: str, 
                             document_type: DocumentType, file_size: int) -> LegalDocument:
        """Create a LegalDocument record."""
        document_id = str(uuid.uuid4())
        
        return LegalDocument(
            id=document_id,
            filename=filename,
            content=content,
            document_type=document_type,
            file_size=file_size
        )
